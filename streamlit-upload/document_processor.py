#!/usr/bin/env python3
"""
Document Processor Module

Handles multi-format document processing including text extraction,
chunking, metadata extraction, and content preprocessing for RAG pipeline.
"""

import os
import json
import csv
import logging
import mimetypes
from typing import List, Dict, Any, Optional, Tuple, Union
from pathlib import Path
from datetime import datetime
from dataclasses import dataclass, asdict
import re

# Import required libraries for document processing
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import markdown
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

# Configure logging
logger = logging.getLogger(__name__)

@dataclass
class DocumentMetadata:
    """Document metadata structure"""
    filename: str
    file_path: str
    file_size: int
    file_type: str
    extension: str
    created_at: str
    processed_at: str
    namespace: str
    page_count: Optional[int] = None
    word_count: Optional[int] = None
    character_count: Optional[int] = None
    language: Optional[str] = None
    title: Optional[str] = None
    author: Optional[str] = None
    subject: Optional[str] = None
    creator: Optional[str] = None
    error_message: Optional[str] = None

@dataclass
class DocumentChunk:
    """Document chunk structure"""
    content: str
    chunk_id: int
    start_char: int
    end_char: int
    word_count: int
    metadata: Dict[str, Any]
    page_number: Optional[int] = None
    section: Optional[str] = None

@dataclass
class ProcessedDocument:
    """Processed document structure"""
    metadata: DocumentMetadata
    chunks: List[DocumentChunk]
    full_text: str
    processing_time: float
    success: bool
    error_message: Optional[str] = None

class DocumentProcessor:
    """
    Multi-format document processor for RAG pipeline
    """
    
    def __init__(self, chunk_size: int = 512, chunk_overlap: int = 50):
        """
        Initialize document processor
        
        Args:
            chunk_size: Size of text chunks in words
            chunk_overlap: Overlap between chunks in words
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.supported_formats = {
            '.pdf': self._process_pdf,
            '.txt': self._process_txt,
            '.md': self._process_markdown,
            '.docx': self._process_docx,
            '.csv': self._process_csv,
            '.json': self._process_json
        }
        
        # Validate availability of required libraries
        self._check_dependencies()
    
    def _check_dependencies(self):
        """Check if required libraries are available"""
        if not PDF_AVAILABLE:
            logger.warning("PyPDF2 not available - PDF processing disabled")
        if not DOCX_AVAILABLE:
            logger.warning("python-docx not available - DOCX processing disabled")
        if not MARKDOWN_AVAILABLE:
            logger.warning("markdown not available - advanced MD processing disabled")
        if not PANDAS_AVAILABLE:
            logger.warning("pandas not available - advanced CSV processing disabled")
    
    def is_supported_format(self, file_path: str) -> bool:
        """
        Check if file format is supported
        
        Args:
            file_path: Path to file
            
        Returns:
            bool: True if format is supported
        """
        extension = Path(file_path).suffix.lower()
        return extension in self.supported_formats
    
    def process_document(self, file_path: str, namespace: str = "default") -> ProcessedDocument:
        """
        Process a document and extract text with chunking
        
        Args:
            file_path: Path to the document file
            namespace: Document namespace for organization
            
        Returns:
            ProcessedDocument: Processed document with metadata and chunks
        """
        start_time = datetime.now()
        
        try:
            # Validate file exists
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"File not found: {file_path}")
            
            # Get file extension
            extension = Path(file_path).suffix.lower()
            
            if extension not in self.supported_formats:
                raise ValueError(f"Unsupported file format: {extension}")
            
            # Create metadata
            metadata = self._create_metadata(file_path, namespace)
            
            # Process file based on format
            processor_func = self.supported_formats[extension]
            full_text, additional_metadata = processor_func(file_path)
            
            # Update metadata with additional info
            metadata.word_count = len(full_text.split())
            metadata.character_count = len(full_text)
            
            # Update with format-specific metadata
            for key, value in additional_metadata.items():
                if hasattr(metadata, key):
                    setattr(metadata, key, value)
            
            # Create chunks
            chunks = self._create_chunks(full_text, metadata)
            
            # Calculate processing time
            processing_time = (datetime.now() - start_time).total_seconds()
            
            return ProcessedDocument(
                metadata=metadata,
                chunks=chunks,
                full_text=full_text,
                processing_time=processing_time,
                success=True
            )
            
        except Exception as e:
            logger.error(f"Error processing document {file_path}: {str(e)}")
            
            # Create error metadata with minimal info
            try:
                metadata = self._create_metadata(file_path, namespace)
            except Exception:
                # If metadata creation fails, create a minimal one
                metadata = DocumentMetadata(
                    filename=Path(file_path).name,
                    file_path=str(file_path),
                    file_size=0,
                    file_type="unknown",
                    extension=Path(file_path).suffix.lower(),
                    created_at=datetime.now().isoformat(),
                    processed_at=datetime.now().isoformat(),
                    namespace=namespace,
                    error_message=str(e)
                )
            
            metadata.error_message = str(e)
            
            processing_time = (datetime.now() - start_time).total_seconds()
            
            return ProcessedDocument(
                metadata=metadata,
                chunks=[],
                full_text="",
                processing_time=processing_time,
                success=False,
                error_message=str(e)
            )
    
    def _create_metadata(self, file_path: str, namespace: str) -> DocumentMetadata:
        """
        Create document metadata
        
        Args:
            file_path: Path to the document
            namespace: Document namespace
            
        Returns:
            DocumentMetadata: Document metadata object
        """
        path_obj = Path(file_path)
        stat = path_obj.stat()
        
        return DocumentMetadata(
            filename=path_obj.name,
            file_path=str(path_obj.absolute()),
            file_size=stat.st_size,
            file_type=mimetypes.guess_type(file_path)[0] or "unknown",
            extension=path_obj.suffix.lower(),
            created_at=datetime.fromtimestamp(stat.st_ctime).isoformat(),
            processed_at=datetime.now().isoformat(),
            namespace=namespace
        )
    
    def _process_pdf(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Process PDF file
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Tuple[str, Dict]: Extracted text and metadata
        """
        if not PDF_AVAILABLE:
            raise ImportError("PyPDF2 not available for PDF processing")
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Extract metadata
                metadata = {}
                if pdf_reader.metadata:
                    metadata['title'] = pdf_reader.metadata.get('/Title', '')
                    metadata['author'] = pdf_reader.metadata.get('/Author', '')
                    metadata['subject'] = pdf_reader.metadata.get('/Subject', '')
                    metadata['creator'] = pdf_reader.metadata.get('/Creator', '')
                
                metadata['page_count'] = len(pdf_reader.pages)
                
                # Extract text from all pages
                text = ""
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text += f"\n--- Page {page_num + 1} ---\n"
                            text += page_text + "\n"
                    except Exception as e:
                        logger.warning(f"Could not extract text from page {page_num + 1}: {e}")
                
                return self._clean_text(text), metadata
                
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {e}")
            raise e
    
    def _process_txt(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Process TXT file
        
        Args:
            file_path: Path to TXT file
            
        Returns:
            Tuple[str, Dict]: Extracted text and metadata
        """
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                    break
                except UnicodeDecodeError:
                    continue
            else:
                raise ValueError("Could not decode file with any supported encoding")
            
            metadata = {}
            lines = text.split('\n')
            
            # Try to extract title from first line if it looks like a header
            if lines and len(lines[0].strip()) < 100 and not lines[0].strip().endswith('.'):
                metadata['title'] = lines[0].strip()
            
            return self._clean_text(text), metadata
            
        except Exception as e:
            logger.error(f"Error processing TXT {file_path}: {e}")
            raise e
    
    def _process_markdown(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Process Markdown file
        
        Args:
            file_path: Path to MD file
            
        Returns:
            Tuple[str, Dict]: Extracted text and metadata
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            metadata = {}
            
            # Extract title from first h1 header
            title_match = re.search(r'^#\s+(.+)', content, re.MULTILINE)
            if title_match:
                metadata['title'] = title_match.group(1).strip()
            
            # Convert markdown to plain text if markdown library is available
            if MARKDOWN_AVAILABLE:
                md = markdown.Markdown()
                html = md.convert(content)
                # Simple HTML tag removal
                text = re.sub(r'<[^>]+>', '', html)
            else:
                # Simple markdown cleanup
                text = content
                text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)  # Remove headers
                text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Remove bold
                text = re.sub(r'\*(.*?)\*', r'\1', text)  # Remove italic
                text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)  # Remove links
            
            return self._clean_text(text), metadata
            
        except Exception as e:
            logger.error(f"Error processing Markdown {file_path}: {e}")
            raise e
    
    def _process_docx(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Process DOCX file
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Tuple[str, Dict]: Extracted text and metadata
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not available for DOCX processing")
        
        try:
            doc = DocxDocument(file_path)
            
            # Extract metadata
            metadata = {}
            core_props = doc.core_properties
            metadata['title'] = core_props.title or ''
            metadata['author'] = core_props.author or ''
            metadata['subject'] = core_props.subject or ''
            metadata['creator'] = core_props.author or ''
            
            # Extract text from paragraphs
            text = ""
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text += " | ".join(row_text) + "\n"
            
            return self._clean_text(text), metadata
            
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {e}")
            raise e
    
    def _process_csv(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Process CSV file
        
        Args:
            file_path: Path to CSV file
            
        Returns:
            Tuple[str, Dict]: Extracted text and metadata
        """
        try:
            metadata = {}
            
            if PANDAS_AVAILABLE:
                # Use pandas for better CSV handling
                df = pd.read_csv(file_path)
                metadata['page_count'] = len(df)  # Number of rows
                
                # Convert to text format
                text = f"CSV Data with {len(df)} rows and {len(df.columns)} columns:\n\n"
                text += f"Columns: {', '.join(df.columns)}\n\n"
                text += df.to_string(index=False)
                
            else:
                # Basic CSV processing
                with open(file_path, 'r', encoding='utf-8') as file:
                    csv_reader = csv.reader(file)
                    rows = list(csv_reader)
                
                if rows:
                    headers = rows[0]
                    metadata['page_count'] = len(rows) - 1
                    
                    text = f"CSV Data with {len(rows)-1} rows and {len(headers)} columns:\n\n"
                    text += f"Columns: {', '.join(headers)}\n\n"
                    
                    for i, row in enumerate(rows[:100]):  # Limit to first 100 rows
                        text += f"Row {i}: {', '.join(row)}\n"
                    
                    if len(rows) > 100:
                        text += f"\n... and {len(rows) - 100} more rows"
                else:
                    text = "Empty CSV file"
            
            return self._clean_text(text), metadata
            
        except Exception as e:
            logger.error(f"Error processing CSV {file_path}: {e}")
            raise e
    
    def _process_json(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Process JSON file
        
        Args:
            file_path: Path to JSON file
            
        Returns:
            Tuple[str, Dict]: Extracted text and metadata
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                data = json.load(file)
            
            metadata = {}
            
            # Convert JSON to readable text
            text = f"JSON Document:\n\n"
            text += self._json_to_text(data, level=0)
            
            # Count items if it's a list or dict
            if isinstance(data, list):
                metadata['page_count'] = len(data)
            elif isinstance(data, dict):
                metadata['page_count'] = len(data)
            
            return self._clean_text(text), metadata
            
        except Exception as e:
            logger.error(f"Error processing JSON {file_path}: {e}")
            raise e
    
    def _json_to_text(self, obj: Any, level: int = 0) -> str:
        """
        Convert JSON object to readable text
        
        Args:
            obj: JSON object
            level: Indentation level
            
        Returns:
            str: Readable text representation
        """
        indent = "  " * level
        text = ""
        
        if isinstance(obj, dict):
            for key, value in obj.items():
                text += f"{indent}{key}:\n"
                if isinstance(value, (dict, list)):
                    text += self._json_to_text(value, level + 1)
                else:
                    text += f"{indent}  {value}\n"
        elif isinstance(obj, list):
            for i, item in enumerate(obj):
                text += f"{indent}Item {i + 1}:\n"
                if isinstance(item, (dict, list)):
                    text += self._json_to_text(item, level + 1)
                else:
                    text += f"{indent}  {item}\n"
        else:
            text += f"{indent}{obj}\n"
        
        return text
    
    def _clean_text(self, text: str) -> str:
        """
        Clean and normalize text
        
        Args:
            text: Raw text
            
        Returns:
            str: Cleaned text
        """
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\n\s*\n', '\n\n', text)  # Multiple newlines to double
        text = re.sub(r'[ \t]+', ' ', text)  # Multiple spaces/tabs to single space
        
        # Remove control characters except newlines and tabs
        text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)
        
        return text.strip()
    
    def _create_chunks(self, text: str, metadata: DocumentMetadata) -> List[DocumentChunk]:
        """
        Create text chunks for embedding
        
        Args:
            text: Full document text
            metadata: Document metadata
            
        Returns:
            List[DocumentChunk]: List of text chunks
        """
        if not text:
            return []
        
        words = text.split()
        chunks = []
        
        start_idx = 0
        chunk_id = 0
        
        while start_idx < len(words):
            # Calculate end index
            end_idx = min(start_idx + self.chunk_size, len(words))
            
            # Get chunk words
            chunk_words = words[start_idx:end_idx]
            chunk_text = ' '.join(chunk_words)
            
            # Calculate character positions (approximate)
            start_char = len(' '.join(words[:start_idx]))
            end_char = start_char + len(chunk_text)
            
            # Create chunk
            chunk = DocumentChunk(
                content=chunk_text,
                chunk_id=chunk_id,
                start_char=start_char,
                end_char=end_char,
                word_count=len(chunk_words),
                metadata={
                    'filename': metadata.filename,
                    'namespace': metadata.namespace,
                    'file_type': metadata.file_type,
                    'chunk_index': chunk_id,
                    'total_chunks': 0  # Will be updated later
                }
            )
            
            chunks.append(chunk)
            
            # Move start index with overlap
            start_idx = max(start_idx + self.chunk_size - self.chunk_overlap, start_idx + 1)
            chunk_id += 1
            
            # Prevent infinite loop
            if start_idx >= len(words):
                break
        
        # Update total chunks count
        for chunk in chunks:
            chunk.metadata['total_chunks'] = len(chunks)
        
        return chunks
    
    def get_processing_stats(self) -> Dict[str, Any]:
        """
        Get processing statistics
        
        Returns:
            Dict: Processing statistics
        """
        return {
            'supported_formats': list(self.supported_formats.keys()),
            'chunk_size': self.chunk_size,
            'chunk_overlap': self.chunk_overlap,
            'dependencies': {
                'pdf_available': PDF_AVAILABLE,
                'docx_available': DOCX_AVAILABLE,
                'markdown_available': MARKDOWN_AVAILABLE,
                'pandas_available': PANDAS_AVAILABLE
            }
        }


# Utility functions for Streamlit integration
def create_document_processor(chunk_size: int = 512, chunk_overlap: int = 50) -> DocumentProcessor:
    """
    Create a document processor instance
    
    Args:
        chunk_size: Size of text chunks in words
        chunk_overlap: Overlap between chunks in words
        
    Returns:
        DocumentProcessor: Configured processor instance
    """
    return DocumentProcessor(chunk_size=chunk_size, chunk_overlap=chunk_overlap)


def process_uploaded_file(file_path: str, namespace: str = "default", 
                         chunk_size: int = 512, chunk_overlap: int = 50) -> ProcessedDocument:
    """
    Process an uploaded file
    
    Args:
        file_path: Path to uploaded file
        namespace: Document namespace
        chunk_size: Chunk size in words
        chunk_overlap: Chunk overlap in words
        
    Returns:
        ProcessedDocument: Processed document result
    """
    processor = DocumentProcessor(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    return processor.process_document(file_path, namespace) 