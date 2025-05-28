#!/usr/bin/env python3
"""
Streamlit Document Upload Application

A comprehensive document upload and processing interface that integrates
with Haystack RAG pipeline, Pinecone, and Redis cache.
"""

import streamlit as st
import os
import sys
import json
import time
from datetime import datetime
from typing import List, Dict, Any, Optional
import logging

# Import document processor and RAG integration
from document_processor import DocumentProcessor, ProcessedDocument, create_document_processor
from rag_integration import RAGPipelineClient, create_rag_client, EmbeddingResult, IndexingResult

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Page configuration
st.set_page_config(
    page_title="Document Upload & Processing",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Constants
SUPPORTED_FORMATS = {
    "PDF": "application/pdf",
    "TXT": "text/plain", 
    "MD": "text/markdown",
    "DOCX": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "CSV": "text/csv",
    "JSON": "application/json"
}

UPLOAD_DIR = "uploads"
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB

# Initialize session state
if 'uploaded_files' not in st.session_state:
    st.session_state.uploaded_files = []
if 'processing_status' not in st.session_state:
    st.session_state.processing_status = {}
if 'namespaces' not in st.session_state:
    st.session_state.namespaces = ["default", "research", "documentation", "general"]

def init_app():
    """Initialize the application"""
    # Create upload directory if it doesn't exist
    os.makedirs(UPLOAD_DIR, exist_ok=True)
    
    # Enhanced CSS for modern UI/UX
    st.markdown("""
    <style>
    .main-header {
        text-align: center;
        padding: 3rem 2rem;
        background: linear-gradient(135deg, #2E8B57 0%, #3CB371 50%, #90EE90 100%);
        color: white;
        border-radius: 20px;
        margin-bottom: 2rem;
        box-shadow: 0 10px 30px rgba(46, 139, 87, 0.3);
        position: relative;
        overflow: hidden;
    }
    
    .main-header::before {
        content: '';
        position: absolute;
        top: -50%;
        left: -50%;
        width: 200%;
        height: 200%;
        background: linear-gradient(45deg, transparent, rgba(255,255,255,0.1), transparent);
        transform: rotate(45deg);
        animation: shine 4s infinite;
    }
    
    @keyframes shine {
        0% { transform: translateX(-100%) translateY(-100%) rotate(45deg); }
        100% { transform: translateX(100%) translateY(100%) rotate(45deg); }
    }
    
    .upload-area {
        border: 3px dashed #2E8B57;
        border-radius: 20px;
        padding: 3rem 2rem;
        text-align: center;
        margin: 2rem 0;
        background: linear-gradient(145deg, #f8f9fa, #e9ecef);
        transition: all 0.3s ease;
        position: relative;
    }
    
    .upload-area:hover {
        border-color: #3CB371;
        background: linear-gradient(145deg, #e9ecef, #f8f9fa);
        transform: translateY(-5px);
        box-shadow: 0 15px 35px rgba(46, 139, 87, 0.2);
    }
    
    .file-info {
        background: linear-gradient(145deg, #e8f5e8, #d4edda);
        padding: 1.5rem;
        border-radius: 15px;
        margin: 1rem 0;
        border-left: 5px solid #2E8B57;
        box-shadow: 0 5px 15px rgba(46, 139, 87, 0.1);
        transition: all 0.3s ease;
        position: relative;
        overflow: hidden;
    }
    
    .file-info:hover {
        transform: translateX(10px);
        box-shadow: 0 8px 25px rgba(46, 139, 87, 0.2);
    }
    
    .file-info::before {
        content: '';
        position: absolute;
        top: 0;
        left: -100%;
        width: 100%;
        height: 100%;
        background: linear-gradient(90deg, transparent, rgba(46, 139, 87, 0.1), transparent);
        transition: left 0.5s;
    }
    
    .file-info:hover::before {
        left: 100%;
    }
    
    .error-message {
        background: linear-gradient(145deg, #ffebee, #f8d7da);
        padding: 1.5rem;
        border-radius: 15px;
        margin: 1rem 0;
        border-left: 5px solid #f44336;
        color: #c62828;
        box-shadow: 0 5px 15px rgba(244, 67, 54, 0.1);
        animation: slideInError 0.5s ease-out;
    }
    
    @keyframes slideInError {
        from { opacity: 0; transform: translateX(-20px); }
        to { opacity: 1; transform: translateX(0); }
    }
    
    .success-message {
        background: linear-gradient(145deg, #e8f5e8, #d4edda);
        padding: 1.5rem;
        border-radius: 15px;
        margin: 1rem 0;
        border-left: 5px solid #2E8B57;
        color: #2e7d32;
        box-shadow: 0 5px 15px rgba(46, 139, 87, 0.1);
        animation: slideInSuccess 0.5s ease-out;
    }
    
    @keyframes slideInSuccess {
        from { opacity: 0; transform: translateY(-20px); }
        to { opacity: 1; transform: translateY(0); }
    }
    
    .info-box {
        background: linear-gradient(145deg, #e3f2fd, #bbdefb);
        padding: 1.5rem;
        border-radius: 15px;
        margin: 1rem 0;
        border-left: 5px solid #2196F3;
        box-shadow: 0 5px 15px rgba(33, 150, 243, 0.1);
        transition: all 0.3s ease;
    }
    
    .info-box:hover {
        transform: translateY(-3px);
        box-shadow: 0 8px 25px rgba(33, 150, 243, 0.2);
    }
    
    .progress-container {
        background: linear-gradient(145deg, #ffffff, #f8f9fa);
        border-radius: 15px;
        padding: 1.5rem;
        margin: 1.5rem 0;
        border: 1px solid #dee2e6;
        box-shadow: 0 5px 15px rgba(0,0,0,0.1);
    }
    
    .processing-animation {
        display: inline-block;
        width: 24px;
        height: 24px;
        border: 3px solid #f3f3f3;
        border-top: 3px solid #2E8B57;
        border-radius: 50%;
        animation: spin 1s linear infinite;
        margin-right: 15px;
        vertical-align: middle;
    }
    
    @keyframes spin {
        0% { transform: rotate(0deg); }
        100% { transform: rotate(360deg); }
    }
    
    .stats-grid {
        display: grid;
        grid-template-columns: repeat(auto-fit, minmax(250px, 1fr));
        gap: 1.5rem;
        margin: 2rem 0;
    }
    
    .stat-card {
        background: linear-gradient(145deg, #ffffff, #f8f9fa);
        border-radius: 20px;
        padding: 2rem;
        text-align: center;
        box-shadow: 0 8px 25px rgba(0,0,0,0.1);
        border-top: 4px solid #2E8B57;
        transition: all 0.3s ease;
        position: relative;
        overflow: hidden;
    }
    
    .stat-card:hover {
        transform: translateY(-8px);
        box-shadow: 0 15px 40px rgba(0,0,0,0.15);
    }
    
    .stat-card::before {
        content: '';
        position: absolute;
        top: 0;
        left: -100%;
        width: 100%;
        height: 100%;
        background: linear-gradient(90deg, transparent, rgba(46, 139, 87, 0.1), transparent);
        transition: left 0.6s;
    }
    
    .stat-card:hover::before {
        left: 100%;
    }
    
    .feature-highlight {
        background: linear-gradient(145deg, #ffffff, #f8f9fa);
        border-radius: 20px;
        padding: 2rem;
        margin: 1.5rem 0;
        box-shadow: 0 8px 25px rgba(0,0,0,0.1);
        border: 1px solid #e9ecef;
        transition: all 0.3s ease;
        position: relative;
    }
    
    .feature-highlight:hover {
        transform: translateY(-5px);
        box-shadow: 0 12px 35px rgba(0,0,0,0.15);
    }
    
    .pulse-animation {
        animation: pulse 2s infinite;
    }
    
    @keyframes pulse {
        0% { transform: scale(1); }
        50% { transform: scale(1.05); }
        100% { transform: scale(1); }
    }
    
    .fade-in {
        animation: fadeIn 0.8s ease-out;
    }
    
    @keyframes fadeIn {
        from { opacity: 0; transform: translateY(20px); }
        to { opacity: 1; transform: translateY(0); }
    }
    
    /* Custom scrollbar */
    ::-webkit-scrollbar {
        width: 8px;
    }
    
    ::-webkit-scrollbar-track {
        background: #f1f1f1;
        border-radius: 10px;
    }
    
    ::-webkit-scrollbar-thumb {
        background: linear-gradient(135deg, #2E8B57, #3CB371);
        border-radius: 10px;
    }
    
    ::-webkit-scrollbar-thumb:hover {
        background: linear-gradient(135deg, #3CB371, #2E8B57);
    }
    </style>
    """, unsafe_allow_html=True)

def validate_file(uploaded_file) -> tuple[bool, str]:
    """
    Validate uploaded file with comprehensive checks
    
    Args:
        uploaded_file: Streamlit uploaded file object
        
    Returns:
        tuple: (is_valid, error_message)
    """
    if uploaded_file is None:
        return False, "No file uploaded"
    
    # Check file size
    if uploaded_file.size == 0:
        return False, "File is empty (0 bytes)"
    
    if uploaded_file.size > MAX_FILE_SIZE:
        return False, f"File size ({uploaded_file.size / (1024*1024):.1f}MB) exceeds maximum allowed size ({MAX_FILE_SIZE / (1024*1024)}MB)"
    
    # Check file extension
    if '.' not in uploaded_file.name:
        return False, "File has no extension"
    
    file_extension = uploaded_file.name.split('.')[-1].upper()
    if file_extension not in SUPPORTED_FORMATS.keys():
        return False, f"File format '.{file_extension.lower()}' is not supported. Supported formats: {', '.join(SUPPORTED_FORMATS.keys())}"
    
    # Check file type
    if uploaded_file.type and uploaded_file.type not in SUPPORTED_FORMATS.values():
        # Some browsers don't set correct MIME types, so we'll be lenient here
        logger.warning(f"MIME type mismatch for {uploaded_file.name}: {uploaded_file.type}")
    
    # Perform content validation based on file type
    try:
        uploaded_file.seek(0)
        content_valid, content_error = validate_file_content(uploaded_file, file_extension.lower())
        if not content_valid:
            return False, content_error
    except Exception as e:
        logger.warning(f"Content validation failed for {uploaded_file.name}: {e}")
        # Don't fail validation for content check errors, just log them
    finally:
        uploaded_file.seek(0)
    
    return True, ""


def validate_file_content(uploaded_file, file_extension: str) -> tuple[bool, str]:
    """
    Validate file content based on file type
    
    Args:
        uploaded_file: Streamlit uploaded file object
        file_extension: File extension (lowercase)
        
    Returns:
        tuple: (is_valid, error_message)
    """
    try:
        if file_extension in ['txt', 'md']:
            # Check if it's valid text
            content = uploaded_file.read()
            try:
                content.decode('utf-8')
            except UnicodeDecodeError:
                try:
                    content.decode('latin-1')
                except UnicodeDecodeError:
                    return False, "File contains invalid text encoding"
        
        elif file_extension == 'json':
            # Check if it's valid JSON
            content = uploaded_file.read().decode('utf-8')
            json.loads(content)
        
        elif file_extension == 'csv':
            # Check if it's valid CSV
            import pandas as pd
            try:
                df = pd.read_csv(uploaded_file)
                if df.empty:
                    return False, "CSV file is empty or has no valid data"
            except Exception:
                # Fallback to basic CSV check
                uploaded_file.seek(0)
                content = uploaded_file.read().decode('utf-8')
                import csv
                csv.reader(content.splitlines())
        
        elif file_extension == 'pdf':
            # Check if it's valid PDF
            try:
                import PyPDF2
                pdf_reader = PyPDF2.PdfReader(uploaded_file)
                if len(pdf_reader.pages) == 0:
                    return False, "PDF file has no pages"
            except ImportError:
                pass  # PyPDF2 not available, skip validation
            except Exception as e:
                return False, f"Invalid PDF file: {str(e)}"
        
        elif file_extension == 'docx':
            # Check if it's valid DOCX
            try:
                from docx import Document
                doc = Document(uploaded_file)
                # Basic validation - just try to access properties
                _ = doc.core_properties
            except ImportError:
                pass  # python-docx not available, skip validation
            except Exception as e:
                return False, f"Invalid DOCX file: {str(e)}"
        
        return True, ""
        
    except json.JSONDecodeError:
        return False, "Invalid JSON format"
    except UnicodeDecodeError:
        return False, "File encoding not supported"
    except Exception as e:
        return False, f"File validation failed: {str(e)}"

def display_file_info(uploaded_file) -> Dict[str, Any]:
    """
    Display file information with preview
    
    Args:
        uploaded_file: Streamlit uploaded file object
        
    Returns:
        dict: File metadata
    """
    file_info = {
        "name": uploaded_file.name,
        "size": uploaded_file.size,
        "type": uploaded_file.type,
        "extension": uploaded_file.name.split('.')[-1].upper(),
        "upload_time": datetime.now().isoformat()
    }
    
    st.markdown(f"""
    <div class="file-info">
        <h4>📄 {file_info['name']}</h4>
        <p><strong>Size:</strong> {file_info['size'] / 1024:.1f} KB</p>
        <p><strong>Type:</strong> {file_info['extension']} ({file_info['type']})</p>
        <p><strong>Upload Time:</strong> {datetime.fromisoformat(file_info['upload_time']).strftime('%Y-%m-%d %H:%M:%S')}</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Add file preview
    display_file_preview(uploaded_file)
    
    return file_info


def display_file_preview(uploaded_file):
    """
    Display file preview based on file type
    
    Args:
        uploaded_file: Streamlit uploaded file object
    """
    file_extension = uploaded_file.name.split('.')[-1].lower()
    
    try:
        # Reset file pointer
        uploaded_file.seek(0)
        
        if file_extension == 'txt':
            preview_text_file(uploaded_file)
        elif file_extension == 'md':
            preview_markdown_file(uploaded_file)
        elif file_extension == 'json':
            preview_json_file(uploaded_file)
        elif file_extension == 'csv':
            preview_csv_file(uploaded_file)
        elif file_extension == 'pdf':
            preview_pdf_file(uploaded_file)
        elif file_extension == 'docx':
            preview_docx_file(uploaded_file)
        else:
            st.info("📄 Preview not available for this file type")
            
    except Exception as e:
        st.warning(f"⚠️ Could not generate preview: {str(e)}")
    finally:
        # Reset file pointer for later use
        uploaded_file.seek(0)


def preview_text_file(uploaded_file):
    """Preview text file content"""
    try:
        content = uploaded_file.read().decode('utf-8')
        preview_content = content[:500] + "..." if len(content) > 500 else content
        
        st.markdown("**📝 Text Preview:**")
        st.text_area("Content", preview_content, height=150, disabled=True)
        
        # Show statistics
        lines = content.count('\n') + 1
        words = len(content.split())
        st.markdown(f"📊 **Stats:** {lines} lines, {words} words, {len(content)} characters")
        
    except UnicodeDecodeError:
        st.warning("⚠️ Could not decode text file - may contain binary data")


def preview_markdown_file(uploaded_file):
    """Preview markdown file content"""
    try:
        content = uploaded_file.read().decode('utf-8')
        preview_content = content[:500] + "..." if len(content) > 500 else content
        
        st.markdown("**📝 Markdown Preview:**")
        
        # Show raw markdown
        with st.expander("Raw Markdown"):
            st.text_area("Raw Content", preview_content, height=150, disabled=True)
        
        # Try to render markdown
        try:
            st.markdown(preview_content)
        except Exception:
            st.text_area("Content", preview_content, height=150, disabled=True)
        
        # Show statistics
        lines = content.count('\n') + 1
        words = len(content.split())
        st.markdown(f"📊 **Stats:** {lines} lines, {words} words, {len(content)} characters")
        
    except UnicodeDecodeError:
        st.warning("⚠️ Could not decode markdown file")


def preview_json_file(uploaded_file):
    """Preview JSON file content"""
    try:
        content = uploaded_file.read().decode('utf-8')
        data = json.loads(content)
        
        st.markdown("**📊 JSON Preview:**")
        
        # Show structure info
        if isinstance(data, dict):
            st.markdown(f"📋 **Type:** Object with {len(data)} keys")
            st.markdown(f"🔑 **Keys:** {', '.join(list(data.keys())[:10])}")
        elif isinstance(data, list):
            st.markdown(f"📋 **Type:** Array with {len(data)} items")
        else:
            st.markdown(f"📋 **Type:** {type(data).__name__}")
        
        # Show formatted JSON (limited)
        preview_data = data
        if isinstance(data, list) and len(data) > 3:
            preview_data = data[:3]
            st.markdown("*Showing first 3 items...*")
        elif isinstance(data, dict) and len(data) > 10:
            preview_data = dict(list(data.items())[:10])
            st.markdown("*Showing first 10 keys...*")
        
        st.json(preview_data)
        
    except json.JSONDecodeError as e:
        st.error(f"❌ Invalid JSON format: {str(e)}")
    except UnicodeDecodeError:
        st.warning("⚠️ Could not decode JSON file")


def preview_csv_file(uploaded_file):
    """Preview CSV file content"""
    try:
        import pandas as pd
        
        # Try to read CSV
        df = pd.read_csv(uploaded_file)
        
        st.markdown("**📊 CSV Preview:**")
        st.markdown(f"📋 **Shape:** {df.shape[0]} rows × {df.shape[1]} columns")
        st.markdown(f"🔑 **Columns:** {', '.join(df.columns.tolist())}")
        
        # Show first few rows
        st.dataframe(df.head(10), use_container_width=True)
        
        if len(df) > 10:
            st.markdown("*Showing first 10 rows...*")
        
        # Show data types
        with st.expander("Column Data Types"):
            st.write(df.dtypes.to_dict())
            
    except Exception as e:
        # Fallback to basic text preview
        uploaded_file.seek(0)
        try:
            content = uploaded_file.read().decode('utf-8')
            lines = content.split('\n')
            
            st.markdown("**📊 CSV Preview (Text Mode):**")
            st.markdown(f"📋 **Lines:** {len(lines)}")
            
            preview_lines = lines[:10]
            st.text_area("Content", '\n'.join(preview_lines), height=150, disabled=True)
            
            if len(lines) > 10:
                st.markdown("*Showing first 10 lines...*")
                
        except Exception:
            st.error(f"❌ Could not preview CSV file: {str(e)}")


def preview_pdf_file(uploaded_file):
    """Preview PDF file content"""
    try:
        import PyPDF2
        
        pdf_reader = PyPDF2.PdfReader(uploaded_file)
        
        st.markdown("**📄 PDF Preview:**")
        st.markdown(f"📋 **Pages:** {len(pdf_reader.pages)}")
        
        # Show metadata if available
        if pdf_reader.metadata:
            metadata = pdf_reader.metadata
            if metadata.get('/Title'):
                st.markdown(f"📑 **Title:** {metadata.get('/Title')}")
            if metadata.get('/Author'):
                st.markdown(f"👤 **Author:** {metadata.get('/Author')}")
            if metadata.get('/Subject'):
                st.markdown(f"📝 **Subject:** {metadata.get('/Subject')}")
        
        # Show first page text
        if len(pdf_reader.pages) > 0:
            try:
                first_page_text = pdf_reader.pages[0].extract_text()
                if first_page_text:
                    preview_text = first_page_text[:500] + "..." if len(first_page_text) > 500 else first_page_text
                    st.text_area("First Page Content", preview_text, height=150, disabled=True)
                else:
                    st.info("📄 Could not extract text from first page")
            except Exception:
                st.info("📄 Could not extract text from PDF")
        
    except ImportError:
        st.warning("⚠️ PyPDF2 not available - PDF preview disabled")
    except Exception as e:
        st.error(f"❌ Could not preview PDF: {str(e)}")


def preview_docx_file(uploaded_file):
    """Preview DOCX file content"""
    try:
        from docx import Document
        
        doc = Document(uploaded_file)
        
        st.markdown("**📄 DOCX Preview:**")
        
        # Show metadata
        core_props = doc.core_properties
        if core_props.title:
            st.markdown(f"📑 **Title:** {core_props.title}")
        if core_props.author:
            st.markdown(f"👤 **Author:** {core_props.author}")
        if core_props.subject:
            st.markdown(f"📝 **Subject:** {core_props.subject}")
        
        # Show paragraph count
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        st.markdown(f"📋 **Paragraphs:** {len(paragraphs)}")
        
        # Show first few paragraphs
        if paragraphs:
            preview_text = '\n\n'.join(paragraphs[:3])
            if len(preview_text) > 500:
                preview_text = preview_text[:500] + "..."
            
            st.text_area("Content Preview", preview_text, height=150, disabled=True)
            
            if len(paragraphs) > 3:
                st.markdown("*Showing first 3 paragraphs...*")
        else:
            st.info("📄 No text content found in document")
        
        # Show table count if any
        if doc.tables:
            st.markdown(f"📊 **Tables:** {len(doc.tables)}")
        
    except ImportError:
        st.warning("⚠️ python-docx not available - DOCX preview disabled")
    except Exception as e:
        st.error(f"❌ Could not preview DOCX: {str(e)}")

def save_uploaded_file(uploaded_file, namespace: str = "default") -> Optional[str]:
    """
    Save uploaded file to disk
    
    Args:
        uploaded_file: Streamlit uploaded file object
        namespace: Organization namespace
        
    Returns:
        str: Saved file path or None if failed
    """
    try:
        # Create namespace directory
        namespace_dir = os.path.join(UPLOAD_DIR, namespace)
        os.makedirs(namespace_dir, exist_ok=True)
        
        # Generate unique filename with timestamp
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{timestamp}_{uploaded_file.name}"
        file_path = os.path.join(namespace_dir, filename)
        
        # Save file
        with open(file_path, "wb") as f:
            f.write(uploaded_file.getbuffer())
        
        logger.info(f"File saved: {file_path}")
        return file_path
        
    except Exception as e:
        logger.error(f"Error saving file: {e}")
        st.error(f"Failed to save file: {str(e)}")
        return None

def render_sidebar():
    """Render sidebar with configuration options"""
    st.sidebar.title("⚙️ Configuration")
    
    # Namespace selection
    st.sidebar.subheader("📁 Namespace")
    selected_namespace = st.sidebar.selectbox(
        "Select or enter namespace:",
        options=st.session_state.namespaces,
        index=0,
        help="Organize your documents by project or category"
    )
    
    # Add custom namespace option
    custom_namespace = st.sidebar.text_input(
        "Or create new namespace:",
        placeholder="my-project",
        help="Enter a new namespace name"
    )
    
    if custom_namespace:
        if custom_namespace not in st.session_state.namespaces:
            st.session_state.namespaces.append(custom_namespace)
            st.rerun()
        selected_namespace = custom_namespace
    
    # Processing options
    st.sidebar.subheader("🔧 Processing Options")
    
    chunk_size = st.sidebar.slider(
        "Chunk Size (words):",
        min_value=100,
        max_value=1000,
        value=512,
        step=50,
        help="Size of text chunks for embedding"
    )
    
    chunk_overlap = st.sidebar.slider(
        "Chunk Overlap (words):",
        min_value=0,
        max_value=200,
        value=50,
        step=10,
        help="Overlap between consecutive chunks"
    )
    
    enable_cache = st.sidebar.checkbox(
        "Enable Redis Cache",
        value=True,
        help="Use Redis cache for faster embedding processing"
    )
    
    # RAG Pipeline options
    st.sidebar.subheader("🚀 RAG Pipeline")
    
    generate_embeddings = st.sidebar.checkbox(
        "Generate Embeddings",
        value=True,
        help="Generate embeddings using Hayhooks RAG pipeline"
    )
    
    index_in_pinecone = st.sidebar.checkbox(
        "Index in Pinecone",
        value=True,
        help="Index embeddings in Pinecone vector store"
    )
    
    hayhooks_url = st.sidebar.text_input(
        "Hayhooks URL:",
                    value="http://hayhooks:1416",
        help="URL of the Hayhooks server"
    )
    
    pipeline_name = st.sidebar.text_input(
        "Pipeline Name:",
        value="rag_pipeline",
        help="Name of the RAG pipeline to use"
    )
    
    # Statistics
    st.sidebar.subheader("📊 Statistics")
    st.sidebar.metric("Files Uploaded", len(st.session_state.uploaded_files))
    
    # Show supported formats
    st.sidebar.subheader("📋 Supported Formats")
    st.sidebar.write(", ".join(SUPPORTED_FORMATS.keys()))
    
    return {
        "namespace": selected_namespace,
        "chunk_size": chunk_size,
        "chunk_overlap": chunk_overlap,
        "enable_cache": enable_cache,
        "generate_embeddings": generate_embeddings,
        "index_in_pinecone": index_in_pinecone,
        "hayhooks_url": hayhooks_url,
        "pipeline_name": pipeline_name
    }

def render_main_interface():
    """Render main upload interface"""
    # Header
    st.markdown("""
    <div class="main-header">
        <h1>📄 Document Upload & Processing</h1>
        <p>Upload documents for embedding and indexing in your RAG system</p>
    </div>
    """, unsafe_allow_html=True)
    
    # File uploader
    st.subheader("📤 Upload Documents")
    
    uploaded_files = st.file_uploader(
        "Choose files to upload",
        type=['pdf', 'txt', 'md', 'docx', 'csv', 'json'],
        accept_multiple_files=True,
        help=f"Supported formats: {', '.join(SUPPORTED_FORMATS.keys())}. Max size: {MAX_FILE_SIZE // (1024*1024)}MB per file."
    )
    
    return uploaded_files

def render_upload_results(uploaded_files, config):
    """Render upload results and file information"""
    if not uploaded_files:
        st.markdown("""
        <div class="info-box">
            <h4>🔍 How to use:</h4>
            <ol>
                <li>Select your namespace in the sidebar</li>
                <li>Configure processing options</li>
                <li>Upload one or more documents</li>
                <li>Preview and validate your files</li>
                <li>Process files to generate embeddings</li>
            </ol>
        </div>
        """, unsafe_allow_html=True)
        return
    
    st.subheader(f"📁 Uploaded Files ({len(uploaded_files)})")
    
    # Process each uploaded file
    valid_files = []
    invalid_files = []
    
    for uploaded_file in uploaded_files:
        st.markdown("---")
        
        # Validate file
        is_valid, error_message = validate_file(uploaded_file)
        
        if is_valid:
            # Display file info
            file_info = display_file_info(uploaded_file)
            file_info['namespace'] = config['namespace']
            file_info['config'] = config
            print(f"🔍 DEBUG: File {uploaded_file.name} assigned to namespace: {config['namespace']}")
            valid_files.append((uploaded_file, file_info))
            
            # Show success message
            st.markdown(f"""
            <div class="success-message">
                ✅ File validation passed - ready for processing
            </div>
            """, unsafe_allow_html=True)
            
        else:
            invalid_files.append((uploaded_file, error_message))
            
            # Show error message
            st.markdown(f"""
            <div class="error-message">
                ❌ {error_message}
            </div>
            """, unsafe_allow_html=True)
    
    # Process valid files
    if valid_files:
        st.markdown("---")
        st.subheader("🔄 Process Files")
        
        col1, col2 = st.columns([2, 1])
        
        with col1:
            if st.button("🚀 RAG Pipeline", type="primary", use_container_width=True):
                process_documents_with_rag(valid_files, config)
        
        with col2:
            st.markdown(f"**Namespace:** `{config['namespace']}`")
    
    # Summary
    if valid_files or invalid_files:
        st.markdown("---")
        st.subheader("📊 Upload Summary")
        
        col1, col2 = st.columns(2)
        with col1:
            st.metric("✅ Valid Files", len(valid_files))
        with col2:
            st.metric("❌ Invalid Files", len(invalid_files))

# Removed: process_valid_files function - functionality integrated into RAG pipeline


# Removed: process_documents_with_chunking function - functionality integrated into RAG pipeline
def _removed_process_documents_with_chunking(valid_files, config):
    """REMOVED: Process documents with enhanced progress tracking and chunking"""
    st.markdown("---")
    
    # Enhanced header with animation
    st.markdown("""
    <div class="feature-highlight fade-in">
        <h3>🧠 Document Processing Pipeline</h3>
        <p>Analyzing and chunking your documents for optimal RAG performance...</p>
    </div>
    """, unsafe_allow_html=True)
    
    if not valid_files:
        st.warning("No valid files to process!")
        return
    
    # First save files if not already saved
    files_to_process = []
    for uploaded_file, file_info in valid_files:
        if 'saved_path' not in file_info:
            saved_path = save_uploaded_file(uploaded_file, file_info['namespace'])
            if saved_path:
                file_info['saved_path'] = saved_path
                files_to_process.append(file_info)
        else:
            files_to_process.append(file_info)
    
    if not files_to_process:
        st.error("No files could be saved for processing!")
        return
    
    # Create document processor
    processor = create_document_processor(
        chunk_size=config['chunk_size'],
        chunk_overlap=config['chunk_overlap']
    )
    
    # Initialize processing stats
    processing_stats = {
        'total_chunks': 0,
        'total_characters': 0,
        'total_words': 0,
        'total_time': 0,
        'successful_files': 0,
        'failed_files': 0
    }
    
    # Enhanced progress tracking
    st.markdown("#### Processing Progress")
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    # Real-time stats container
    stats_container = st.container()
    
    # Results container
    results_container = st.container()
    processing_results = []
    
    for i, file_info in enumerate(files_to_process):
        current_progress = (i + 1) / len(files_to_process)
        
        # Enhanced status display
        with status_text:
            st.markdown(f"""
            <div class="progress-container fade-in">
                <div class="processing-animation"></div>
                <strong>Processing {file_info['name']}...</strong> ({i+1}/{len(files_to_process)})
                <br><small>Extracting text and creating chunks...</small>
            </div>
            """, unsafe_allow_html=True)
        
        # Update real-time stats
        with stats_container:
            col1, col2, col3, col4 = st.columns(4)
            with col1:
                st.markdown(f"""
                <div class="stat-card">
                    <h4>{processing_stats['successful_files']}</h4>
                    <p>Files Processed</p>
                </div>
                """, unsafe_allow_html=True)
            with col2:
                st.markdown(f"""
                <div class="stat-card">
                    <h4>{processing_stats['total_chunks']}</h4>
                    <p>Chunks Created</p>
                </div>
                """, unsafe_allow_html=True)
            with col3:
                st.markdown(f"""
                <div class="stat-card">
                    <h4>{processing_stats['total_words']:,}</h4>
                    <p>Words Processed</p>
                </div>
                """, unsafe_allow_html=True)
            with col4:
                st.markdown(f"""
                <div class="stat-card">
                    <h4>{processing_stats['total_time']:.1f}s</h4>
                    <p>Processing Time</p>
                </div>
                """, unsafe_allow_html=True)
        
        try:
            # Process document
            result = processor.process_document(
                file_path=file_info['saved_path'],
                namespace=file_info['namespace']
            )
            processing_results.append(result)
            
            # Update processing stats
            if result.success:
                processing_stats['successful_files'] += 1
                processing_stats['total_chunks'] += len(result.chunks)
                processing_stats['total_characters'] += len(result.full_text)
                processing_stats['total_words'] += len(result.full_text.split())
                processing_stats['total_time'] += result.processing_time
                
                # Show success in results container
                with results_container:
                    st.markdown(f"""
                    <div class="success-message fade-in">
                        <h4>✅ {file_info['name']}</h4>
                        <div class="stats-grid">
                            <div><strong>Chunks:</strong> {len(result.chunks)}</div>
                            <div><strong>Words:</strong> {len(result.full_text.split()):,}</div>
                            <div><strong>Time:</strong> {result.processing_time:.2f}s</div>
                        </div>
                    </div>
                    """, unsafe_allow_html=True)
            else:
                processing_stats['failed_files'] += 1
                with results_container:
                    st.markdown(f"""
                    <div class="error-message fade-in">
                        <h4>❌ {file_info['name']}</h4>
                        <p><strong>Error:</strong> {result.error_message}</p>
                    </div>
                    """, unsafe_allow_html=True)
            
            # Update progress
            progress_bar.progress(current_progress)
            time.sleep(0.2)  # Smooth animation
            
        except Exception as e:
            processing_stats['failed_files'] += 1
            with results_container:
                st.markdown(f"""
                <div class="error-message fade-in">
                    <h4>❌ {file_info['name']}</h4>
                    <p><strong>Error:</strong> {str(e)}</p>
                </div>
                """, unsafe_allow_html=True)
            logger.error(f"Error processing {file_info['name']}: {e}")
    
    # Final status with celebration
    status_text.markdown("""
    <div class="success-message pulse-animation">
        <h3>🎉 Processing Complete!</h3>
        <p>All documents have been processed and are ready for use.</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Display enhanced results
    if processing_results:
        successful_results = [r for r in processing_results if r.success]
        failed_results = [r for r in processing_results if not r.success]
        
        # Enhanced summary with detailed metrics
        with st.expander("📊 Detailed Processing Summary", expanded=True):
            col1, col2 = st.columns(2)
            
            with col1:
                st.markdown(f"""
                <div class="feature-highlight">
                    <h4>📈 Processing Statistics</h4>
                    <ul>
                        <li><strong>Total Files:</strong> {len(files_to_process)}</li>
                        <li><strong>Successful:</strong> {processing_stats['successful_files']}</li>
                        <li><strong>Failed:</strong> {processing_stats['failed_files']}</li>
                        <li><strong>Success Rate:</strong> {(processing_stats['successful_files']/len(files_to_process)*100):.1f}%</li>
                    </ul>
                </div>
                """, unsafe_allow_html=True)
            
            with col2:
                st.markdown(f"""
                <div class="feature-highlight">
                    <h4>📊 Content Analysis</h4>
                    <ul>
                        <li><strong>Total Chunks:</strong> {processing_stats['total_chunks']:,}</li>
                        <li><strong>Total Words:</strong> {processing_stats['total_words']:,}</li>
                        <li><strong>Total Characters:</strong> {processing_stats['total_characters']:,}</li>
                        <li><strong>Processing Time:</strong> {processing_stats['total_time']:.2f}s</li>
                    </ul>
                </div>
                """, unsafe_allow_html=True)
            
            # Performance metrics
            if processing_stats['total_time'] > 0:
                words_per_second = processing_stats['total_words'] / processing_stats['total_time']
                chunks_per_second = processing_stats['total_chunks'] / processing_stats['total_time']
                
                st.markdown(f"""
                <div class="info-box">
                    <h4>⚡ Performance Metrics</h4>
                    <p><strong>Processing Speed:</strong> {words_per_second:.0f} words/second</p>
                    <p><strong>Chunking Rate:</strong> {chunks_per_second:.1f} chunks/second</p>
                    <p><strong>Average Chunk Size:</strong> {processing_stats['total_words']/max(processing_stats['total_chunks'], 1):.0f} words</p>
                </div>
                """, unsafe_allow_html=True)
        
        # Show detailed results for each document
        if successful_results:
            with st.expander("📄 Document Details", expanded=False):
                for result in successful_results:
                    st.markdown(f"### 📄 {result.metadata.filename}")
                    
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        st.markdown("**📊 Document Info:**")
                        st.write(f"• **Size:** {result.metadata.file_size / 1024:.1f} KB")
                        st.write(f"• **Words:** {result.metadata.word_count:,}")
                        st.write(f"• **Characters:** {result.metadata.character_count:,}")
                        st.write(f"• **Chunks:** {len(result.chunks)}")
                        st.write(f"• **Processing Time:** {result.processing_time:.2f}s")
                        
                        if result.metadata.title:
                            st.write(f"• **Title:** {result.metadata.title}")
                        if result.metadata.author:
                            st.write(f"• **Author:** {result.metadata.author}")
                    
                    with col2:
                        st.markdown("**📝 Text Preview:**")
                        preview_text = result.full_text[:300] + "..." if len(result.full_text) > 300 else result.full_text
                        st.text_area("Document Content", preview_text, height=150, disabled=True, key=f"preview_{result.metadata.filename}")
                    
                    # Show chunk preview
                    if result.chunks:
                        st.markdown("**🧩 Chunks Preview:**")
                        for idx, chunk in enumerate(result.chunks[:3]):  # Show first 3 chunks
                            # Handle both string chunks and DocumentChunk objects
                            chunk_text = chunk.content if hasattr(chunk, 'content') else str(chunk)
                            preview_text = chunk_text[:200] + '...' if len(chunk_text) > 200 else chunk_text
                            st.markdown(f"""
                            <div class="file-info">
                                <strong>Chunk {idx + 1}:</strong><br>
                                {preview_text}
                            </div>
                            """, unsafe_allow_html=True)
                        if len(result.chunks) > 3:
                            st.info(f"... and {len(result.chunks) - 3} more chunks")
                    
                    st.markdown("---")
        
        # Store results in session state for future use
        if 'processed_documents' not in st.session_state:
            st.session_state.processed_documents = []
        
        st.session_state.processed_documents.extend(processing_results)
        
        # Final success message
        if successful_results:
            st.markdown("""
            <div class="success-message pulse-animation">
                <h4>🚀 Ready for Next Steps!</h4>
                <p>Documents are processed and ready for embedding generation and indexing in your RAG pipeline.</p>
            </div>
            """, unsafe_allow_html=True)
    else:
        st.warning("No documents were processed successfully.")


def process_documents_with_rag(valid_files, config):
    """Process documents with full RAG pipeline (chunking + embeddings + indexing)"""
    st.markdown("---")
    
    # Enhanced header with animation
    st.markdown("""
    <div class="feature-highlight fade-in">
        <h3>🚀 RAG Pipeline Processing</h3>
        <p>Processing documents through the complete RAG pipeline: chunking → embeddings → indexing</p>
    </div>
    """, unsafe_allow_html=True)
    
    if not valid_files:
        st.warning("No valid files to process!")
        return
    
    # Check RAG configuration
    if not config.get('generate_embeddings', False) and not config.get('index_in_pinecone', False):
        st.warning("⚠️ Please enable at least one RAG option: Generate Embeddings or Index in Pinecone")
        return
    
    # Initialize RAG client
    try:
        rag_client = create_rag_client(
            hayhooks_url=config.get('hayhooks_url', 'http://hayhooks:1416'),
            pipeline_name=config.get('pipeline_name', 'rag_pipeline'),
            use_cache=config.get('enable_cache', True)
        )
        
        # Health check
        st.info("🔍 Checking RAG pipeline connection...")
        if not rag_client.health_check():
            st.error(f"❌ Cannot connect to Hayhooks server at {config.get('hayhooks_url')}. Please check the server status.")
            return
        
        st.success(f"✅ Connected to RAG pipeline: {config.get('pipeline_name')}")
        
    except Exception as e:
        st.error(f"❌ Failed to initialize RAG client: {str(e)}")
        return
    
    # First save and process files if not already done
    files_to_process = []
    for uploaded_file, file_info in valid_files:
        if 'saved_path' not in file_info:
            saved_path = save_uploaded_file(uploaded_file, file_info['namespace'])
            if saved_path:
                file_info['saved_path'] = saved_path
                files_to_process.append(file_info)
        else:
            files_to_process.append(file_info)
    
    if not files_to_process:
        st.error("No files could be saved for processing!")
        return
    
    # Create document processor
    processor = create_document_processor(
        chunk_size=config['chunk_size'],
        chunk_overlap=config['chunk_overlap']
    )
    
    # Initialize processing stats
    processing_stats = {
        'total_chunks': 0,
        'total_embeddings': 0,
        'total_indexed': 0,
        'total_time': 0,
        'successful_files': 0,
        'failed_files': 0,
        'embedding_time': 0,
        'indexing_time': 0
    }
    
    # Enhanced progress tracking
    st.markdown("#### RAG Pipeline Progress")
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    # Real-time stats container
    stats_container = st.container()
    
    # Results container
    results_container = st.container()
    rag_results = []
    
    for i, file_info in enumerate(files_to_process):
        current_progress = (i + 1) / len(files_to_process)
        
        # Enhanced status display
        with status_text:
            st.markdown(f"""
            <div class="progress-container fade-in">
                <div class="processing-animation"></div>
                <strong>Processing {file_info['name']} through RAG pipeline...</strong> ({i+1}/{len(files_to_process)})
                <br><small>Step 1/3: Document processing and chunking...</small>
            </div>
            """, unsafe_allow_html=True)
        
        try:
            # Step 1: Process document and create chunks
            doc_result = processor.process_document(
                file_path=file_info['saved_path'],
                namespace=file_info['namespace']
            )
            
            if not doc_result.success:
                processing_stats['failed_files'] += 1
                with results_container:
                    st.markdown(f"""
                    <div class="error-message fade-in">
                        <h4>❌ {file_info['name']} - Document Processing Failed</h4>
                        <p><strong>Error:</strong> {doc_result.error_message}</p>
                    </div>
                    """, unsafe_allow_html=True)
                continue
            
            # Prepare metadata for each chunk
            metadata_list = []
            for idx, chunk in enumerate(doc_result.chunks):
                # Handle both string chunks and DocumentChunk objects
                chunk_text = chunk.content if hasattr(chunk, 'content') else str(chunk)
                metadata = {
                    "document_id": doc_result.metadata.filename.replace('.', '_'),
                    "chunk_id": f"chunk_{idx}",
                    "filename": doc_result.metadata.filename,
                    "namespace": file_info['namespace'],
                    "chunk_index": idx,
                    "total_chunks": len(doc_result.chunks),
                    "file_size": doc_result.metadata.file_size,
                    "word_count": len(chunk_text.split()),
                    "timestamp": datetime.now().isoformat()
                }
                
                # Add document metadata if available
                if doc_result.metadata.title:
                    metadata["title"] = doc_result.metadata.title
                if doc_result.metadata.author:
                    metadata["author"] = doc_result.metadata.author
                
                metadata_list.append(metadata)
            
            embedding_results = []
            indexing_result = None
            
            # Step 2: Generate embeddings if enabled
            if config.get('generate_embeddings', False):
                with status_text:
                    st.markdown(f"""
                    <div class="progress-container fade-in">
                        <div class="processing-animation"></div>
                        <strong>Processing {file_info['name']} through RAG pipeline...</strong> ({i+1}/{len(files_to_process)})
                        <br><small>Step 2/3: Generating embeddings ({len(doc_result.chunks)} chunks)...</small>
                    </div>
                    """, unsafe_allow_html=True)
                
                # Extract text from chunks (handle both string and DocumentChunk objects)
                chunk_texts = []
                for chunk in doc_result.chunks:
                    if hasattr(chunk, 'content'):
                        chunk_texts.append(chunk.content)
                    else:
                        chunk_texts.append(str(chunk))
                
                start_time = time.time()
                embedding_results = rag_client.generate_embeddings(
                    chunks=chunk_texts,
                    metadata_list=metadata_list,
                    namespace=file_info['namespace']
                )
                processing_stats['embedding_time'] += time.time() - start_time
                
                successful_embeddings = [r for r in embedding_results if r.success]
                processing_stats['total_embeddings'] += len(successful_embeddings)
                
                # Step 3: Index in Pinecone if enabled
                if config.get('index_in_pinecone', False) and successful_embeddings:
                    with status_text:
                        st.markdown(f"""
                        <div class="progress-container fade-in">
                            <div class="processing-animation"></div>
                            <strong>Processing {file_info['name']} through RAG pipeline...</strong> ({i+1}/{len(files_to_process)})
                            <br><small>Step 3/3: Indexing in Pinecone ({len(successful_embeddings)} embeddings)...</small>
                        </div>
                        """, unsafe_allow_html=True)
                    
                    start_time = time.time()
                    indexing_result = rag_client.index_document(
                        embedding_results=successful_embeddings,
                        namespace=file_info['namespace']
                    )
                    processing_stats['indexing_time'] += time.time() - start_time
                    
                    if indexing_result.success:
                        processing_stats['total_indexed'] += indexing_result.indexed_chunks
            
            # Update processing stats
            processing_stats['successful_files'] += 1
            processing_stats['total_chunks'] += len(doc_result.chunks)
            processing_stats['total_time'] += doc_result.processing_time
            
            # Store results
            rag_result = {
                'document': doc_result,
                'embeddings': embedding_results,
                'indexing': indexing_result,
                'file_info': file_info
            }
            rag_results.append(rag_result)
            
            # Show success in results container
            with results_container:
                embedding_count = len([r for r in embedding_results if r.success]) if embedding_results else 0
                indexed_count = indexing_result.indexed_chunks if indexing_result and indexing_result.success else 0
                
                st.markdown(f"""
                <div class="success-message fade-in">
                    <h4>✅ {file_info['name']} - RAG Pipeline Complete</h4>
                    <div class="stats-grid">
                        <div><strong>Chunks:</strong> {len(doc_result.chunks)}</div>
                        <div><strong>Embeddings:</strong> {embedding_count}</div>
                        <div><strong>Indexed:</strong> {indexed_count}</div>
                        <div><strong>Time:</strong> {doc_result.processing_time:.2f}s</div>
                    </div>
                </div>
                """, unsafe_allow_html=True)
            
        except Exception as e:
            processing_stats['failed_files'] += 1
            with results_container:
                st.markdown(f"""
                <div class="error-message fade-in">
                    <h4>❌ {file_info['name']} - RAG Pipeline Error</h4>
                    <p><strong>Error:</strong> {str(e)}</p>
                </div>
                """, unsafe_allow_html=True)
            logger.error(f"Error in RAG pipeline for {file_info['name']}: {e}")
        
        # Update real-time stats
        with stats_container:
            col1, col2, col3, col4 = st.columns(4)
            with col1:
                st.markdown(f"""
                <div class="stat-card">
                    <h4>{processing_stats['successful_files']}</h4>
                    <p>Files Processed</p>
                </div>
                """, unsafe_allow_html=True)
            with col2:
                st.markdown(f"""
                <div class="stat-card">
                    <h4>{processing_stats['total_embeddings']}</h4>
                    <p>Embeddings Generated</p>
                </div>
                """, unsafe_allow_html=True)
            with col3:
                st.markdown(f"""
                <div class="stat-card">
                    <h4>{processing_stats['total_indexed']}</h4>
                    <p>Vectors Indexed</p>
                </div>
                """, unsafe_allow_html=True)
            with col4:
                st.markdown(f"""
                <div class="stat-card">
                    <h4>{processing_stats['total_time']:.1f}s</h4>
                    <p>Total Time</p>
                </div>
                """, unsafe_allow_html=True)
        
        # Update progress
        progress_bar.progress(current_progress)
        time.sleep(0.2)  # Smooth animation
    
    # Final status with celebration
    status_text.markdown("""
    <div class="success-message pulse-animation">
        <h3>🎉 RAG Pipeline Complete!</h3>
        <p>All documents have been processed through the RAG pipeline and are ready for querying.</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Display enhanced results
    if rag_results:
        successful_results = [r for r in rag_results if r['document'].success]
        
        # Enhanced summary with detailed metrics
        with st.expander("📊 RAG Pipeline Summary", expanded=True):
            col1, col2 = st.columns(2)
            
            with col1:
                st.markdown(f"""
                <div class="feature-highlight">
                    <h4>📈 Pipeline Statistics</h4>
                    <ul>
                        <li><strong>Total Files:</strong> {len(files_to_process)}</li>
                        <li><strong>Successful:</strong> {processing_stats['successful_files']}</li>
                        <li><strong>Failed:</strong> {processing_stats['failed_files']}</li>
                        <li><strong>Success Rate:</strong> {(processing_stats['successful_files']/len(files_to_process)*100):.1f}%</li>
                    </ul>
                </div>
                """, unsafe_allow_html=True)
            
            with col2:
                st.markdown(f"""
                <div class="feature-highlight">
                    <h4>🚀 RAG Performance</h4>
                    <ul>
                        <li><strong>Total Chunks:</strong> {processing_stats['total_chunks']:,}</li>
                        <li><strong>Embeddings Generated:</strong> {processing_stats['total_embeddings']:,}</li>
                        <li><strong>Vectors Indexed:</strong> {processing_stats['total_indexed']:,}</li>
                        <li><strong>Embedding Time:</strong> {processing_stats['embedding_time']:.2f}s</li>
                        <li><strong>Indexing Time:</strong> {processing_stats['indexing_time']:.2f}s</li>
                    </ul>
                </div>
                """, unsafe_allow_html=True)
        
        # Show detailed results for each document
        if successful_results:
            with st.expander("📄 Document RAG Results", expanded=False):
                for result in successful_results:
                    doc = result['document']
                    embeddings = result['embeddings']
                    indexing = result['indexing']
                    
                    st.markdown(f"### 🚀 {doc.metadata.filename}")
                    
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        st.markdown("**📊 Document Processing:**")
                        st.write(f"• **Chunks Created:** {len(doc.chunks)}")
                        st.write(f"• **Words:** {doc.metadata.word_count:,}")
                        st.write(f"• **Processing Time:** {doc.processing_time:.2f}s")
                        
                        if embeddings:
                            successful_embeddings = [e for e in embeddings if e.success]
                            st.write(f"• **Embeddings Generated:** {len(successful_embeddings)}/{len(embeddings)}")
                            
                            if successful_embeddings:
                                embedding_dim = len(successful_embeddings[0].embedding)
                                st.write(f"• **Embedding Dimensions:** {embedding_dim}")
                    
                    with col2:
                        st.markdown("**🗃️ Indexing Results:**")
                        if indexing:
                            st.write(f"• **Indexed Chunks:** {indexing.indexed_chunks}")
                            st.write(f"• **Failed Chunks:** {indexing.failed_chunks}")
                            st.write(f"• **Indexing Time:** {indexing.processing_time:.2f}s")
                            if indexing.pinecone_ids:
                                st.write(f"• **Pinecone IDs:** {len(indexing.pinecone_ids)} vectors")
                        else:
                            st.write("• **Status:** Indexing not performed")
                    
                    st.markdown("---")
        
        # Store results in session state for future use
        if 'rag_processed_documents' not in st.session_state:
            st.session_state.rag_processed_documents = []
        
        st.session_state.rag_processed_documents.extend(rag_results)
        
        # Final success message
        if successful_results:
            st.markdown("""
            <div class="success-message pulse-animation">
                <h4>🎯 RAG Pipeline Ready!</h4>
                <p>Documents are fully processed and indexed. Your knowledge base is ready for semantic search and retrieval.</p>
            </div>
            """, unsafe_allow_html=True)
    else:
        st.warning("No documents were processed through the RAG pipeline successfully.")


def render_dashboard():
    """Render enhanced dashboard with session statistics"""
    if 'processed_documents' in st.session_state and st.session_state.processed_documents:
        st.markdown("---")
        
        # Dashboard header
        st.markdown("""
        <div class="feature-highlight fade-in">
            <h3>📊 Session Dashboard</h3>
            <p>Overview of your document processing session</p>
        </div>
        """, unsafe_allow_html=True)
        
        # Calculate session stats
        processed_docs = st.session_state.processed_documents
        successful_docs = [doc for doc in processed_docs if doc.success]
        
        total_chunks = sum(len(doc.chunks) for doc in successful_docs)
        total_words = sum(len(doc.full_text.split()) for doc in successful_docs)
        total_chars = sum(len(doc.full_text) for doc in successful_docs)
        total_time = sum(doc.processing_time for doc in processed_docs)
        
        # Enhanced metrics display
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            st.markdown(f"""
            <div class="stat-card pulse-animation">
                <h2>{len(successful_docs)}</h2>
                <p>Documents Processed</p>
                <small>{len(processed_docs) - len(successful_docs)} failed</small>
            </div>
            """, unsafe_allow_html=True)
        
        with col2:
            st.markdown(f"""
            <div class="stat-card pulse-animation">
                <h2>{total_chunks:,}</h2>
                <p>Total Chunks</p>
                <small>Ready for embedding</small>
            </div>
            """, unsafe_allow_html=True)
        
        with col3:
            st.markdown(f"""
            <div class="stat-card pulse-animation">
                <h2>{total_words:,}</h2>
                <p>Words Processed</p>
                <small>{total_chars:,} characters</small>
            </div>
            """, unsafe_allow_html=True)
        
        with col4:
            avg_chunk_size = total_words / max(total_chunks, 1)
            st.markdown(f"""
            <div class="stat-card pulse-animation">
                <h2>{avg_chunk_size:.0f}</h2>
                <p>Avg Chunk Size</p>
                <small>words per chunk</small>
            </div>
            """, unsafe_allow_html=True)
        
        # Performance insights
        if total_time > 0:
            processing_speed = total_words / total_time
            
            with st.expander("⚡ Performance Insights", expanded=False):
                col1, col2 = st.columns(2)
                
                with col1:
                    st.markdown(f"""
                    <div class="info-box">
                        <h4>🚀 Processing Speed</h4>
                        <p><strong>{processing_speed:.0f}</strong> words/second</p>
                        <p><strong>{total_chunks/total_time:.1f}</strong> chunks/second</p>
                        <p><strong>{total_time:.2f}</strong> seconds total</p>
                    </div>
                    """, unsafe_allow_html=True)
                
                with col2:
                    # Document type breakdown
                    doc_types = {}
                    for doc in successful_docs:
                        ext = doc.metadata.filename.split('.')[-1].upper()
                        doc_types[ext] = doc_types.get(ext, 0) + 1
                    
                    st.markdown("""
                    <div class="info-box">
                        <h4>📄 Document Types</h4>
                    """, unsafe_allow_html=True)
                    
                    for doc_type, count in doc_types.items():
                        st.markdown(f"<p><strong>{doc_type}:</strong> {count} files</p>", unsafe_allow_html=True)
                    
                    st.markdown("</div>", unsafe_allow_html=True)
        
        # Quick actions
        st.markdown("### 🚀 Quick Actions")
        col1, col2, col3 = st.columns(3)
        
        with col1:
            if st.button("🗑️ Clear Session", type="secondary"):
                st.session_state.processed_documents = []
                st.session_state.uploaded_files = []
                st.rerun()
        
        with col2:
            if st.button("📊 Export Stats", type="secondary"):
                stats_data = {
                    "session_summary": {
                        "total_documents": len(processed_docs),
                        "successful_documents": len(successful_docs),
                        "total_chunks": total_chunks,
                        "total_words": total_words,
                        "total_characters": total_chars,
                        "processing_time": total_time,
                        "average_chunk_size": avg_chunk_size,
                        "processing_speed": processing_speed if total_time > 0 else 0
                    },
                    "documents": [
                        {
                            "filename": doc.metadata.filename,
                            "success": doc.success,
                            "chunks": len(doc.chunks) if doc.success else 0,
                            "words": len(doc.full_text.split()) if doc.success else 0,
                            "processing_time": doc.processing_time
                        }
                        for doc in processed_docs
                    ]
                }
                
                st.download_button(
                    label="📥 Download JSON",
                    data=json.dumps(stats_data, indent=2),
                    file_name=f"processing_stats_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                    mime="application/json"
                )
        
        with col3:
            if st.button("🔄 Process More", type="primary"):
                st.info("Upload more files above to continue processing!")

    # RAG Query Interface
    if 'rag_processed_documents' in st.session_state and st.session_state.rag_processed_documents:
        render_rag_query_interface()

def render_rag_query_interface():
    """Render interface for querying RAG-processed documents"""
    st.markdown("---")
    
    # Header
    st.markdown("""
    <div class="feature-highlight fade-in">
        <h3>🔍 Semantic Search</h3>
        <p>Query your processed documents using semantic search</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Get available namespaces from processed documents
    rag_docs = st.session_state.rag_processed_documents
    available_namespaces = list(set([doc['file_info']['namespace'] for doc in rag_docs]))
    
    col1, col2 = st.columns([3, 1])
    
    with col1:
        query_text = st.text_input(
            "Enter your query:",
            placeholder="What is the main topic discussed in the documents?",
            help="Enter a question or topic to search for in your indexed documents"
        )
    
    with col2:
        query_namespace = st.selectbox(
            "Namespace:",
            options=["all"] + available_namespaces,
            help="Select namespace to search in"
        )
    
    col1, col2, col3 = st.columns([1, 1, 2])
    
    with col1:
        top_k = st.slider("Results:", min_value=1, max_value=10, value=5)
    
    with col2:
        if st.button("🔍 Search", type="primary"):
            if query_text:
                perform_semantic_search(query_text, query_namespace, top_k)
            else:
                st.warning("Please enter a search query")
    
    with col3:
        st.markdown(f"**Available Documents:** {len(rag_docs)} processed")

def perform_semantic_search(query_text: str, namespace: str, top_k: int):
    """Perform semantic search on RAG-processed documents"""
    try:
        # Get RAG client configuration from session state or use defaults
        config = {
            'hayhooks_url': 'http://hayhooks:1416',
            'pipeline_name': 'rag_pipeline',
            'enable_cache': True
        }
        
        # Initialize RAG client
        rag_client = create_rag_client(
            hayhooks_url=config['hayhooks_url'],
            pipeline_name=config['pipeline_name'],
            use_cache=config['enable_cache']
        )
        
        # Check connection
        if not rag_client.health_check():
            st.error(f"❌ Cannot connect to Hayhooks server at {config['hayhooks_url']}")
            return
        
        # Perform search
        with st.spinner("🔍 Searching through your documents..."):
            # Use the selected namespace or find the most common one from processed docs
            if namespace == "all":
                # Get the most common namespace from processed documents
                rag_docs = st.session_state.rag_processed_documents
                namespace_counts = {}
                for doc in rag_docs:
                    ns = doc['file_info']['namespace']
                    namespace_counts[ns] = namespace_counts.get(ns, 0) + 1
                search_namespace = max(namespace_counts.keys(), key=lambda k: namespace_counts[k]) if namespace_counts else "default"
            else:
                search_namespace = namespace
            
            results = rag_client.query_similar_documents(
                query_text=query_text,
                namespace=search_namespace,
                top_k=top_k
            )
        
        if results:
            st.markdown("### 📋 Search Results")
            
            for i, result in enumerate(results, 1):
                # Extract metadata
                metadata = result.get('metadata', {})
                filename = metadata.get('filename', 'Unknown')
                chunk_index = metadata.get('chunk_index', 0)
                similarity_score = result.get('score', 0.0)
                content = result.get('content', result.get('text', ''))
                
                # Display result
                with st.expander(f"🔍 Result {i}: {filename} (Chunk {chunk_index + 1})", expanded=i==1):
                    col1, col2 = st.columns([2, 1])
                    
                    with col1:
                        st.markdown("**📝 Content:**")
                        st.markdown(f"*{content[:500]}{'...' if len(content) > 500 else ''}*")
                    
                    with col2:
                        st.markdown("**📊 Metadata:**")
                        st.write(f"**File:** {filename}")
                        st.write(f"**Chunk:** {chunk_index + 1}")
                        st.write(f"**Similarity:** {similarity_score:.4f}")
                        
                        if metadata.get('namespace'):
                            st.write(f"**Namespace:** {metadata['namespace']}")
                        if metadata.get('word_count'):
                            st.write(f"**Words:** {metadata['word_count']}")
                        if metadata.get('title'):
                            st.write(f"**Title:** {metadata['title']}")
                
                st.markdown("---")
        else:
            st.warning("🔍 No results found for your query. Try different keywords or check if documents are properly indexed.")
            
    except Exception as e:
        st.error(f"❌ Search failed: {str(e)}")
        logger.error(f"Semantic search error: {e}")

def main():
    """Main application function"""
    # Initialize app
    init_app()
    
    # Render sidebar and get configuration
    config = render_sidebar()
    
    # Render main interface
    uploaded_files = render_main_interface()
    
    # Render upload results
    render_upload_results(uploaded_files, config)
    
    # Render dashboard if there are processed documents
    render_dashboard()
    
    # Enhanced footer with system status
    st.markdown("---")
    
    # System status indicators
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.markdown("""
        <div class="info-box">
            <h4>🔧 System Status</h4>
            <p>✅ Streamlit Frontend</p>
            <p>✅ Document Processor</p>
            <p>✅ File Validation</p>
        </div>
        """, unsafe_allow_html=True)
    
    with col2:
        st.markdown("""
        <div class="info-box">
            <h4>📊 Supported Formats</h4>
            <p>PDF, TXT, MD, DOCX</p>
            <p>CSV, JSON</p>
            <p>Max: 50MB per file</p>
        </div>
        """, unsafe_allow_html=True)
    
    with col3:
        st.markdown("""
        <div class="info-box">
            <h4>🚀 Integration Ready</h4>
            <p>Haystack RAG Pipeline</p>
            <p>Pinecone Vector Store</p>
            <p>Redis Caching</p>
        </div>
        """, unsafe_allow_html=True)
    
    # Powered by footer
    st.markdown("""
    <div style="text-align: center; color: #666; padding: 2rem; margin-top: 2rem; border-top: 1px solid #eee;">
        <h4>🚀 Powered by Modern AI Stack</h4>
        <p>Streamlit • Haystack • Pinecone • Redis • FastAPI • Docker</p>
        <small>Built for high-performance document processing and RAG applications</small>
    </div>
    """, unsafe_allow_html=True)

if __name__ == "__main__":
    main() 